import io
import re
import docx

# Add this line to ignore invalid XML characters in the docx module
docx.oxml.ns.qn('w:eastAsianLayout')

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from pdfminer.high_level import extract_pages
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.pdfpage import PDFPage

# Define the section titles to search for
section_titles = ["Drawings", "Claim Objections", "Claim Rejections", "Conclusion"]

# Create a dictionary to store the sections
sections = {title: "" for title in section_titles}

# Open the PDF file
with open("officeaction.pdf", "rb") as f:
    # Create a PDF resource manager and set parameters
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()

    # Create a TextConverter object to extract text from the PDF file
    output = io.BytesIO()
    device = TextConverter(rsrcmgr, outfp=output, laparams=laparams)

    # Create a PDF page interpreter object to process each page in the PDF file
    interpreter = PDFPageInterpreter(rsrcmgr, device)

    # Process each page in the PDF file
    for page in PDFPage.get_pages(f):
        interpreter.process_page(page)

    # Get the text from the TextConverter output
    text = output.getvalue().decode()

    # Search for each section title in the text
    for title in section_titles:
        # Find the start and end positions of the section
        pattern = re.compile(rf"{title}\b", re.IGNORECASE)
        matches = [(m.start(), m.end()) for m in pattern.finditer(text)]

        # If the section was found, extract the text and add it to the sections dictionary
        if matches:
            start = matches[0][1]
            if title == "Conclusion":
                end = len(text)
            else:
                end = matches[-1][0]
            sections[title] += text[start:end] + "\n"
            print(f"{title}: {sections[title]}")

# Generate the PDF report
pdf_buffer = io.BytesIO()
pdf_doc = canvas.Canvas(pdf_buffer, pagesize=letter)

# Set the font size and leading for the report
font_size = 14
leading = font_size * 1.5

# Set the initial y-coordinate for the report
y = letter[1] - 50

# Loop through each section and add it to the report
for title, section in sections.items():
    # Print the section title
    pdf_doc.setFont("Helvetica-Bold", font_size)
    pdf_doc.drawString(50, y, title)
    y -= leading

    # Print the section text
    pdf_doc.setFont("Helvetica", font_size)
    lines = section.split("\n")
    for line in lines:
        pdf_doc.drawString(50, y, line)
        y -= leading

    # Add some space between sections
    y -= leading

# Save the PDF report to a file
with open("pdfsections.pdf", "wb") as f:
    f.write(pdf_buffer.getvalue())

# Generate the Word report
docx_doc = docx.Document()

# Loop through each section and add it to the document
for title, section in sections.items():
    # Add the section title to the document
    docx_doc.add_heading(title, level=1)

    # Add the section text to the document
    try:
        docx_doc.add_paragraph(section)
    except ValueError:
        # Ignore any errors related to invalid XML characters
        pass

# Save the Word document to a file
docx_doc.save("docxsections.docx")

# Print the sections
for title, section in sections.items():
    print(f"{title}:\n{section}\n")

